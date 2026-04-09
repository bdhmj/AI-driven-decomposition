"""Telegram bot for AI-driven project estimation and decomposition."""

import asyncio
import io
import logging
import os
import re
import tempfile
from urllib.parse import urljoin, urlparse

from datetime import date, datetime, timedelta

import httpx
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv
from openpyxl import Workbook
from openpyxl.drawing.image import Image as XlImage
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from openpyxl.utils import get_column_letter
from telegram import InlineKeyboardButton, InlineKeyboardMarkup, Update
from telegram.constants import ParseMode
from telegram.ext import (
    Application,
    CallbackQueryHandler,
    CommandHandler,
    ConversationHandler,
    MessageHandler,
    filters,
)

import claude_service

load_dotenv(override=True)

logging.basicConfig(
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    level=logging.INFO,
)
logger = logging.getLogger(__name__)

# Conversation states
(
    WAITING_REQUEST,
    WAITING_ANSWERS,
    REVIEWING_SPEC,
    WAITING_SPEC_FEEDBACK,
    REVIEWING_DECOMPOSITION,
    WAITING_COEFFICIENTS,
    WAITING_RATES,
    WAITING_MARGIN,
    PROCESSING,
) = range(9)

# ─── Default coefficients and rates ─────────────────────────────────────────

DEFAULT_COEFFICIENTS = {
    "debug_pct": 20,          # Проверка и отладка задач (%)
    "code_review_hours": 1,   # Код ревью (часов/день на разработчика)
    "communication_hours": 3, # Коммуникации (часов/неделю на каждого)
    "qa_pct": 0,              # Тестировщик (% от общего, 0 если QA оценён задачами)
    "risk_buffer_pct": 20,    # Буфер на риски (%)
    "devops_pct": 0,          # DevOps доп. (%)
    "pm_pct": 30,             # Менеджер (% от самого длинного специалиста)
}

DEFAULT_RATES = {
    "DevOps": 25,
    "Smart contract": 30,
    "Backend": 25,
    "Frontend": 20,
    "QA": 15,
    "Manual QA": 15,
    "UX/UI дизайнер": 20,
    "Аналитик": 20,
    "Mobile Developer": 25,
    "Data Engineer": 30,
}

COEFF_LABELS = {
    "debug_pct": "Проверка и отладка задач (%)",
    "code_review_hours": "Код ревью (часов/день)",
    "communication_hours": "Коммуникации (часов/неделю)",
    "qa_pct": "Тестировщик (% от общего)",
    "risk_buffer_pct": "Буфер на риски (%)",
    "devops_pct": "DevOps доп. (%)",
    "pm_pct": "Менеджер (% от макс. специалиста)",
}


def calc_K(coeffs: dict) -> float:
    """Calculate the project coefficient K from PM parameters."""
    return (
        1
        + coeffs["code_review_hours"] / 8
        + coeffs["communication_hours"] / 40
        + coeffs["debug_pct"] / 100
        + coeffs["risk_buffer_pct"] / 100
        + coeffs["devops_pct"] / 100
    )

# Per-user session storage
sessions: dict[int, dict] = {}

claude_client = None


def get_session(user_id: int) -> dict:
    if user_id not in sessions:
        sessions[user_id] = {}
    return sessions[user_id]


def escape_md(text: str) -> str:
    """Escape special characters for MarkdownV2."""
    special = r"_*[]()~`>#+-=|{}.!\\"
    result = []
    for ch in text:
        if ch in special:
            result.append("\\")
        result.append(ch)
    return "".join(result)


# ─── File extraction helpers ─────────────────────────────────────────────────

SKIP_EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".svg", ".ico",
                   ".mp3", ".mp4", ".avi", ".mov", ".zip", ".rar", ".gz",
                   ".tar", ".7z", ".exe", ".dll", ".woff", ".woff2", ".ttf",
                   ".eot", ".css", ".js", ".map"}


def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """Extract text from uploaded file based on extension."""
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".txt" or ext == ".md" or ext == ".csv" or ext == ".log":
        for encoding in ("utf-8", "cp1251", "latin-1"):
            try:
                return file_bytes.decode(encoding)
            except UnicodeDecodeError:
                continue
        return file_bytes.decode("utf-8", errors="replace")

    if ext == ".docx":
        doc = Document(io.BytesIO(file_bytes))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    if ext == ".doc":
        # Fallback: try reading as raw text
        return file_bytes.decode("utf-8", errors="replace")

    if ext == ".pdf":
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                pages = [p.extract_text() or "" for p in pdf.pages]
                return "\n\n".join(pages)
        except ImportError:
            return "[PDF поддержка: установите pdfplumber]"

    if ext == ".rtf":
        try:
            from striprtf.striprtf import rtf_to_text
            return rtf_to_text(file_bytes.decode("utf-8", errors="replace"))
        except ImportError:
            return file_bytes.decode("utf-8", errors="replace")

    if ext == ".odt":
        try:
            from odf.opendocument import load as odf_load
            from odf.text import P as OdfP
            from odf import teletype
            doc = odf_load(io.BytesIO(file_bytes))
            paragraphs = doc.getElementsByType(OdfP)
            return "\n".join(teletype.extractText(p) for p in paragraphs)
        except ImportError:
            return "[ODT поддержка: установите odfpy]"

    # Try as plain text
    return file_bytes.decode("utf-8", errors="replace")


# ─── Google Docs helper ─────────────────────────────────────────────────────

GDOC_PATTERNS = [
    re.compile(r"docs\.google\.com/document/d/([a-zA-Z0-9_-]+)"),
]


def extract_gdoc_id(url: str) -> str | None:
    for pat in GDOC_PATTERNS:
        m = pat.search(url)
        if m:
            return m.group(1)
    return None


async def fetch_gdoc_text(url: str) -> str | None:
    doc_id = extract_gdoc_id(url)
    if not doc_id:
        return None
    export_url = f"https://docs.google.com/document/d/{doc_id}/export?format=txt"
    async with httpx.AsyncClient(follow_redirects=True, timeout=15) as client:
        resp = await client.get(export_url)
        if resp.status_code == 200:
            return resp.text
    return None


# ─── Website crawler ─────────────────────────────────────────────────────────

def _extract_page_text(html: str) -> str:
    """Extract meaningful text from HTML, stripping scripts/styles."""
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "noscript", "svg", "iframe"]):
        tag.decompose()
    text = soup.get_text(separator="\n", strip=True)
    # Collapse excessive blank lines
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    return "\n".join(lines)


def _extract_links(html: str, base_url: str, base_domain: str) -> set[str]:
    """Extract internal links from HTML."""
    soup = BeautifulSoup(html, "html.parser")
    links = set()
    for a in soup.find_all("a", href=True):
        href = a["href"].split("#")[0].split("?")[0]  # Remove fragments and query
        if not href or href.startswith(("mailto:", "tel:", "javascript:")):
            continue
        full_url = urljoin(base_url, href)
        parsed = urlparse(full_url)
        if parsed.netloc != base_domain:
            continue
        ext = os.path.splitext(parsed.path)[1].lower()
        if ext in SKIP_EXTENSIONS:
            continue
        # Normalize
        clean = f"{parsed.scheme}://{parsed.netloc}{parsed.path}"
        if clean.endswith("/"):
            clean = clean[:-1]
        links.add(clean)
    return links


async def crawl_website(url: str, max_pages: int = 30) -> str:
    """Crawl a website, visiting internal pages. Returns combined text."""
    parsed = urlparse(url)
    base_domain = parsed.netloc
    if not base_domain:
        return ""

    visited = set()
    to_visit = {url.split("#")[0].split("?")[0]}
    pages_text = []

    # Also try to fetch sitemap
    sitemap_url = f"{parsed.scheme}://{base_domain}/sitemap.xml"

    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
    }

    async with httpx.AsyncClient(follow_redirects=True, timeout=10, headers=headers) as client:
        # Try sitemap first
        try:
            resp = await client.get(sitemap_url)
            if resp.status_code == 200 and "<loc>" in resp.text:
                soup = BeautifulSoup(resp.text, "html.parser")
                for loc in soup.find_all("loc"):
                    loc_url = loc.text.strip()
                    loc_parsed = urlparse(loc_url)
                    if loc_parsed.netloc == base_domain:
                        ext = os.path.splitext(loc_parsed.path)[1].lower()
                        if ext not in SKIP_EXTENSIONS:
                            to_visit.add(loc_url.split("#")[0].split("?")[0])
        except Exception:
            pass

        while to_visit and len(visited) < max_pages:
            current = to_visit.pop()
            if current in visited:
                continue
            visited.add(current)

            try:
                resp = await client.get(current)
                if resp.status_code != 200:
                    continue
                content_type = resp.headers.get("content-type", "")
                if "text/html" not in content_type:
                    continue
                html = resp.text
            except Exception:
                continue

            page_text = _extract_page_text(html)
            if page_text:
                pages_text.append(f"=== СТРАНИЦА: {current} ===\n{page_text}")

            # Extract more links
            new_links = _extract_links(html, current, base_domain)
            for link in new_links:
                if link not in visited:
                    to_visit.add(link)

    return "\n\n".join(pages_text)


# ─── URL detection ───────────────────────────────────────────────────────────

URL_REGEX = re.compile(r'https?://[^\s<>"\']+')


def find_urls(text: str) -> list[str]:
    return URL_REGEX.findall(text)


# ─── Handlers ────────────────────────────────────────────────────────────────


async def cmd_start(update: Update, context) -> None:
    await update.message.reply_text(
        "Привет! Я бот для первичной оценки клиентских запросов на разработку.\n\n"
        "Используй /new_request чтобы начать оценку нового проекта."
    )


async def cmd_new_request(update: Update, context) -> int:
    user_id = update.effective_user.id
    sessions[user_id] = {}
    await update.message.reply_text(
        "Пришли мне клиентский запрос на разработку.\n\n"
        "Поддерживаемые форматы:\n"
        "• Текстовое сообщение\n"
        "• Файл (.txt, .docx, .pdf, .rtf, .odt, .md, .csv)\n"
        "• Ссылка на Google Документ\n"
        "• Ссылка на сайт-референс (бот проанализирует все разделы)"
    )
    return WAITING_REQUEST


async def _extract_request_text(update: Update) -> tuple[str | None, bool]:
    """Extract request text from message. Returns (text, is_website_url).
    If is_website_url=True, text contains the raw URLs + user note (crawling not done yet)."""
    msg = update.message

    # 1) Document/file upload
    if msg.document:
        file_obj = await msg.document.get_file()
        file_bytes = await file_obj.download_as_bytearray()
        filename = msg.document.file_name or "file.txt"
        text = extract_text_from_file(bytes(file_bytes), filename)
        if msg.caption:
            text = msg.caption + "\n\n" + text
        return text, False

    # 2) Text message — check for links
    if msg.text:
        text = msg.text
        urls = find_urls(text)

        if urls:
            # Check for Google Docs link first
            for url in urls:
                gdoc_text = await fetch_gdoc_text(url)
                if gdoc_text:
                    user_note = URL_REGEX.sub("", text).strip()
                    result = gdoc_text
                    if user_note:
                        result = user_note + "\n\n" + result
                    return result, False

            # Website URL — return for async processing
            return text, True

        # Plain text (no URLs)
        return text, False

    return None, False


async def _run_crawl_and_analyze(chat_id: int, user_id: int, urls: list[str], user_note: str, context) -> None:
    """Background task: crawl website, summarize, analyze, send results."""
    session = get_session(user_id)

    try:
        # Crawl
        all_pages = []
        for url in urls:
            if session.get("_cancelled"):
                return
            pages = await crawl_website(url, max_pages=30)
            if pages:
                all_pages.append(pages)

        if session.get("_cancelled"):
            return

        if not all_pages:
            await context.bot.send_message(chat_id, "Не удалось получить данные с сайта. Попробуй /new_request")
            return

        combined = "\n\n".join(all_pages)
        if len(combined) > 100000:
            combined = combined[:100000]

        await context.bot.send_message(chat_id, "✅ Сайт просканирован. Анализирую содержимое через AI...")

        if session.get("_cancelled"):
            return

        # Summarize website via Claude
        project_desc = await asyncio.to_thread(
            claude_service.summarize_website, claude_client, combined
        )

        if user_note:
            project_desc = user_note + "\n\n" + project_desc

        session["request"] = project_desc

        if session.get("_cancelled"):
            return

        # Analyze request
        await context.bot.send_message(chat_id, "⏳ Анализирую запрос...")
        analysis = await asyncio.to_thread(
            claude_service.analyze_request, claude_client, project_desc
        )

        if session.get("_cancelled"):
            return

        if analysis.get("sufficient"):
            # Generate spec
            await context.bot.send_message(chat_id, "⏳ Составляю техническое задание...")
            spec = await asyncio.to_thread(
                claude_service.generate_spec, claude_client, project_desc
            )
            session["spec"] = spec
            doc_bytes = _build_spec_docx(spec)
            await context.bot.send_document(
                chat_id, document=doc_bytes,
                filename="Техническое_задание.docx",
                caption="📄 Техническое задание",
            )
            keyboard = InlineKeyboardMarkup([
                [
                    InlineKeyboardButton("✅ Декомпозировать", callback_data="decompose"),
                    InlineKeyboardButton("🔄 Заново", callback_data="redo_spec"),
                ]
            ])
            await context.bot.send_message(
                chat_id, "Проверь ТЗ в файле выше. Что делаем дальше?",
                reply_markup=keyboard,
            )
            session["_proc_next_state"] = REVIEWING_SPEC
        else:
            questions = analysis.get("questions", [])
            session["questions"] = questions
            text = "Для качественной оценки мне нужна дополнительная информация:\n\n"
            for i, q in enumerate(questions, 1):
                text += f"{i}. {q}\n\n"
            text += "Пришли ответы одним сообщением или нажми кнопку ниже."
            keyboard = InlineKeyboardMarkup(
                [[InlineKeyboardButton("Ответов нет, далее ➡️", callback_data="skip_answers")]]
            )
            await context.bot.send_message(chat_id, text, reply_markup=keyboard)
            session["_proc_next_state"] = WAITING_ANSWERS

    except asyncio.CancelledError:
        return
    except Exception as e:
        logger.error("Error in crawl_and_analyze: %s", e)
        await context.bot.send_message(
            chat_id, f"Ошибка при анализе сайта: {e}\nПопробуй /new_request"
        )


async def handle_request(update: Update, context) -> int:
    """Stage 2: Receive client request — text, file, Google Doc, or URL."""
    user_id = update.effective_user.id
    session = get_session(user_id)

    request_text, is_website = await _extract_request_text(update)

    if is_website:
        # Website URL — start background crawl, enter PROCESSING state
        urls = find_urls(request_text)
        user_note = URL_REGEX.sub("", request_text).strip()
        await update.message.reply_text(
            "🔍 Анализирую сайт — обхожу все разделы, это может занять 1-2 минуты...\n"
            "Для отмены: /skip"
        )
        task = asyncio.create_task(
            _run_crawl_and_analyze(update.effective_chat.id, user_id, urls, user_note, context)
        )
        session["_task"] = task
        session["_cancelled"] = False
        return PROCESSING

    if not request_text or not request_text.strip():
        await update.message.reply_text(
            "Не удалось извлечь текст из сообщения. Попробуй отправить ещё раз в другом формате."
        )
        return WAITING_REQUEST

    session["request"] = request_text

    await update.message.reply_text("⏳ Анализирую запрос...")

    try:
        analysis = claude_service.analyze_request(claude_client, request_text)
    except Exception as e:
        logger.error("Claude API error in analyze_request: %s", e)
        await update.message.reply_text(
            "Произошла ошибка при анализе запроса. Попробуй ещё раз /new_request"
        )
        return ConversationHandler.END

    if analysis.get("sufficient"):
        return await _generate_spec(update, session)
    else:
        questions = analysis.get("questions", [])
        session["questions"] = questions
        text = "Для качественной оценки мне нужна дополнительная информация. Пожалуйста, ответь на следующие вопросы (передай их клиенту при необходимости):\n\n"
        for i, q in enumerate(questions, 1):
            text += f"{i}. {q}\n\n"
        text += "Пришли ответы одним сообщением или нажми кнопку ниже."
        keyboard = InlineKeyboardMarkup(
            [[InlineKeyboardButton("Ответов нет, далее ➡️", callback_data="skip_answers")]]
        )
        await update.message.reply_text(text, reply_markup=keyboard)
        return WAITING_ANSWERS


async def handle_skip_answers(update: Update, context) -> int:
    """Stage 2b: Skip answers and proceed to spec generation."""
    query = update.callback_query
    await query.answer()
    user_id = query.from_user.id
    session = get_session(user_id)
    # Use query.message as the "update.message" for _generate_spec
    return await _generate_spec_from_msg(query.message, session)


async def handle_answers(update: Update, context) -> int:
    """Stage 2b: Receive answers to clarifying questions."""
    user_id = update.effective_user.id
    session = get_session(user_id)
    session["answers"] = update.message.text
    return await _generate_spec(update, session)


async def _generate_spec_impl(message, session: dict) -> int:
    """Stage 3: Generate spec, send as docx."""
    await message.reply_text("⏳ Составляю техническое задание...")

    try:
        spec = claude_service.generate_spec(
            claude_client,
            session["request"],
            answers=session.get("answers"),
            feedback=session.get("spec_feedback"),
            previous_spec=session.get("spec"),
        )
    except Exception as e:
        logger.error("Claude API error in generate_spec: %s", e)
        await message.reply_text("Ошибка при генерации ТЗ. Попробуй /new_request")
        return ConversationHandler.END

    session["spec"] = spec

    doc_bytes = _build_spec_docx(spec)
    await message.reply_document(
        document=doc_bytes,
        filename="Техническое_задание.docx",
        caption="📄 Техническое задание",
    )

    keyboard = InlineKeyboardMarkup(
        [
            [
                InlineKeyboardButton("✅ Декомпозировать", callback_data="decompose"),
                InlineKeyboardButton("🔄 Заново", callback_data="redo_spec"),
            ]
        ]
    )
    await message.reply_text(
        "Проверь ТЗ в файле выше. Что делаем дальше?", reply_markup=keyboard
    )
    return REVIEWING_SPEC


async def _generate_spec_from_msg(message, session: dict) -> int:
    """Stage 3: Generate spec from callback query message."""
    return await _generate_spec_impl(message, session)


async def _generate_spec(update: Update, session: dict) -> int:
    """Stage 3: Generate technical specification."""
    return await _generate_spec_impl(update.message, session)


def _build_spec_docx(spec_text: str) -> io.BytesIO:
    """Convert spec text (with markdown-like headings) into a formatted .docx."""
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)

    for line in spec_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue

        # Detect markdown headings
        heading_match = re.match(r"^(#{1,4})\s+(.*)", stripped)
        if heading_match:
            level = min(len(heading_match.group(1)), 4)
            text = heading_match.group(2).strip("*").strip()
            doc.add_heading(text, level=level)
            continue

        # Detect numbered section titles (e.g. "1. Title" or "1.2. Title")
        section_match = re.match(r"^(\d+\.[\d.]*)\s+\*\*(.*?)\*\*", stripped)
        if section_match:
            doc.add_heading(f"{section_match.group(1)} {section_match.group(2)}", level=2)
            continue

        # Bold lines (full line wrapped in **)
        if stripped.startswith("**") and stripped.endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(stripped.strip("*").strip())
            run.bold = True
            continue

        # Bullet points
        if stripped.startswith("- ") or stripped.startswith("• "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
            continue

        # Regular paragraph — strip inline bold markers for clean text
        clean = re.sub(r"\*\*(.*?)\*\*", r"\1", stripped)
        doc.add_paragraph(clean)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    buf.name = "spec.docx"
    return buf


async def handle_spec_review(update: Update, context) -> int:
    """Stage 4: Handle spec review buttons."""
    query = update.callback_query
    await query.answer()
    user_id = query.from_user.id
    session = get_session(user_id)

    if query.data == "decompose":
        return await _decompose(query, session)
    elif query.data == "redo_spec":
        await query.message.reply_text(
            "Опиши, что не так в текущем ТЗ. Какие замечания нужно учесть?"
        )
        return WAITING_SPEC_FEEDBACK


async def handle_spec_feedback(update: Update, context) -> int:
    """Stage 4b: Receive feedback and regenerate spec."""
    user_id = update.effective_user.id
    session = get_session(user_id)
    session["spec_feedback"] = update.message.text
    return await _generate_spec(update, session)


async def _decompose(query, session: dict) -> int:
    """Stage 5: Decompose spec into task modules with min/max day estimates."""
    await query.message.reply_text("⏳ Декомпозирую задачи и оцениваю трудозатраты...")

    try:
        data = claude_service.decompose_tasks(claude_client, session["spec"])
    except Exception as e:
        logger.error("Claude API error in decompose_tasks: %s", e)
        await query.message.reply_text(
            "Ошибка при декомпозиции. Попробуй /new_request"
        )
        return ConversationHandler.END

    session["modules"] = data["modules"]

    # Flatten tasks for later use
    all_tasks = []
    for m in data["modules"]:
        for t in m["tasks"]:
            all_tasks.append(t)
    session["all_tasks"] = all_tasks

    # Show decomposition grouped by modules
    lines = ["📋 ДЕКОМПОЗИЦИЯ ЗАДАЧ (дни min / max)", "=" * 60, ""]
    for m in data["modules"]:
        lines.append(f"📁 {m['name']}")
        for t in m["tasks"]:
            comment = f" — {t['comment']}" if t.get("comment") else ""
            lines.append(
                f"  • {t['task']} [{t['specialist']}] ({t['min_days']}-{t['max_days']} дн){comment}"
            )
        lines.append("")

    # Summary: total days per specialist
    spec_days: dict[str, dict] = {}
    for t in all_tasks:
        name = t["specialist"]
        if name not in spec_days:
            spec_days[name] = {"min": 0, "max": 0}
        spec_days[name]["min"] += t["min_days"]
        spec_days[name]["max"] += t["max_days"]

    lines.append("👥 СВОДКА ПО СПЕЦИАЛИСТАМ (дни без коэффициентов)")
    lines.append("-" * 40)
    for name, d in spec_days.items():
        lines.append(f"  • {name}: {d['min']}-{d['max']} дней")
    lines.append("")

    text = "\n".join(lines)
    for chunk in _split_message(text):
        await query.message.reply_text(chunk)

    session["spec_days"] = spec_days

    # Now ask for coefficients
    return await _show_coefficients(query.message, session)


async def _show_coefficients(message, session: dict) -> int:
    """Show current coefficients and ask admin to adjust."""
    coeffs = session.get("coefficients", DEFAULT_COEFFICIENTS.copy())
    session["coefficients"] = coeffs
    K = calc_K(coeffs)

    lines = ["⚙️ КОЭФФИЦИЕНТЫ ПРОЕКТА", "=" * 40, ""]
    keys = list(COEFF_LABELS.keys())
    for i, key in enumerate(keys, 1):
        lines.append(f"{i}. {COEFF_LABELS[key]}: {coeffs[key]}")
    lines.append("")
    lines.append(f"📐 Итого K = {K:.2f}")
    lines.append("")
    lines.append("Чтобы изменить параметр — отправь номер и новое значение через пробел.")
    lines.append("Например: 5 25 (буфер на риски → 25%)")

    keyboard = InlineKeyboardMarkup(
        [[InlineKeyboardButton("✅ Далее — к ставкам", callback_data="coeffs_done")]]
    )
    await message.reply_text("\n".join(lines), reply_markup=keyboard)
    return WAITING_COEFFICIENTS


async def handle_coefficients(update: Update, context) -> int:
    """Handle coefficient adjustment or confirm."""
    user_id = update.effective_user.id
    session = get_session(user_id)
    text = update.message.text.strip()

    parts = text.split(None, 1)
    if len(parts) != 2:
        await update.message.reply_text("Формат: номер значение. Например: 5 25")
        return WAITING_COEFFICIENTS

    try:
        idx = int(parts[0])
        val = float(parts[1].replace(",", "."))
    except ValueError:
        await update.message.reply_text("Номер и значение должны быть числами.")
        return WAITING_COEFFICIENTS

    keys = list(COEFF_LABELS.keys())
    if idx < 1 or idx > len(keys):
        await update.message.reply_text(f"Номер должен быть от 1 до {len(keys)}.")
        return WAITING_COEFFICIENTS

    session["coefficients"][keys[idx - 1]] = val
    return await _show_coefficients(update.message, session)


async def handle_coefficients_done(update: Update, context) -> int:
    """Coefficients confirmed — move to rates."""
    query = update.callback_query
    await query.answer()
    user_id = query.from_user.id
    session = get_session(user_id)
    return await _show_rates(query.message, session)


async def _show_rates(message, session: dict) -> int:
    """Show internal rates for each specialist on this project."""
    spec_days = session["spec_days"]
    coeffs = session["coefficients"]

    # Build rates dict for specialists on this project
    if "rates" not in session:
        rates = {}
        for name in spec_days:
            rates[name] = DEFAULT_RATES.get(name, 25)
        session["rates"] = rates
    rates = session["rates"]

    lines = ["💰 ВНУТРЕННИЕ СТАВКИ ($/ч на руки)", "=" * 40, ""]
    for name in spec_days:
        lines.append(f"  • {name}: ${rates.get(name, 25)}/ч")
    lines.append("")

    # Also show PM rate if PM is auto-calculated
    if coeffs["pm_pct"] > 0:
        pm_rate = rates.get("Project manager", DEFAULT_RATES.get("Project manager", 15))
        rates["Project manager"] = pm_rate
        lines.append(f"  • Project manager: ${pm_rate}/ч (авто {coeffs['pm_pct']}% от макс. специалиста)")

    if coeffs["qa_pct"] > 0 and "QA (авто)" not in rates:
        qa_rate = rates.get("QA", DEFAULT_RATES.get("QA", 15))
        rates["QA (авто)"] = qa_rate
        lines.append(f"  • QA (авто): ${qa_rate}/ч ({coeffs['qa_pct']}% от общего)")

    lines.append("")
    lines.append("Чтобы изменить ставку — отправь: Название_специалиста Ставка")
    lines.append("Например: Frontend 22")

    keyboard = InlineKeyboardMarkup(
        [[InlineKeyboardButton("✅ Далее — к марже", callback_data="rates_done")]]
    )
    await message.reply_text("\n".join(lines), reply_markup=keyboard)
    return WAITING_RATES


async def handle_rates(update: Update, context) -> int:
    """Handle internal rate adjustment."""
    user_id = update.effective_user.id
    session = get_session(user_id)
    text = update.message.text.strip()

    parts = text.rsplit(" ", 1)
    if len(parts) != 2:
        await update.message.reply_text("Формат: НазваниеСпециалиста Ставка\nНапример: Frontend 22")
        return WAITING_RATES

    name, rate_str = parts
    try:
        new_rate = float(rate_str)
    except ValueError:
        await update.message.reply_text("Ставка должна быть числом.")
        return WAITING_RATES

    # Find specialist (case-insensitive)
    found = False
    for key in list(session["rates"].keys()):
        if key.lower() == name.lower():
            session["rates"][key] = new_rate
            found = True
            break

    if not found:
        names = ", ".join(session["rates"].keys())
        await update.message.reply_text(f"Специалист '{name}' не найден. Доступные: {names}")
        return WAITING_RATES

    return await _show_rates(update.message, session)


async def handle_rates_done(update: Update, context) -> int:
    """Rates confirmed — ask for margin."""
    query = update.callback_query
    await query.answer()
    await query.message.reply_text(
        "Введи маржу агентства в процентах (единая для всех специалистов).\n"
        "Например: 50"
    )
    return WAITING_MARGIN


async def handle_margin(update: Update, context) -> int:
    """Receive margin and produce final report."""
    user_id = update.effective_user.id
    session = get_session(user_id)
    text = update.message.text.strip().replace("%", "").replace(",", ".")

    try:
        margin = float(text)
    except ValueError:
        await update.message.reply_text("Введи число (процент маржи). Например: 50")
        return WAITING_MARGIN

    session["margin"] = margin
    coeffs = session["coefficients"]
    rates = session["rates"]
    spec_days = session["spec_days"]
    K = calc_K(coeffs)

    # Calculate final weeks and costs per specialist
    specialists = []
    max_weeks = 0
    total_weeks_all = 0

    for name, d in spec_days.items():
        avg_days = (d["min"] + d["max"]) / 2
        final_days = avg_days * K
        weeks = final_days / 5
        hours = final_days * 8
        rate = rates.get(name, 25)
        cost = hours * rate
        client_rate = rate * (1 + margin / 100)
        client_cost = hours * client_rate

        specialists.append({
            "name": name,
            "weeks": round(weeks, 2),
            "days": round(final_days, 1),
            "hours": round(hours, 1),
            "rate": rate,
            "cost": round(cost, 2),
            "client_rate": round(client_rate, 2),
            "client_cost": round(client_cost, 2),
        })
        max_weeks = max(max_weeks, weeks)
        total_weeks_all += weeks

    # Add PM if pm_pct > 0
    if coeffs["pm_pct"] > 0:
        pm_weeks = max_weeks * coeffs["pm_pct"] / 100
        pm_hours = pm_weeks * 5 * 8
        pm_rate = rates.get("Project manager", 15)
        pm_client_rate = pm_rate * (1 + margin / 100)
        specialists.append({
            "name": "Project manager",
            "weeks": round(pm_weeks, 2),
            "days": round(pm_weeks * 5, 1),
            "hours": round(pm_hours, 1),
            "rate": pm_rate,
            "cost": round(pm_hours * pm_rate, 2),
            "client_rate": round(pm_client_rate, 2),
            "client_cost": round(pm_hours * pm_client_rate, 2),
        })

    # Add auto-QA if qa_pct > 0
    if coeffs["qa_pct"] > 0:
        qa_weeks = total_weeks_all * coeffs["qa_pct"] / 100
        qa_hours = qa_weeks * 5 * 8
        qa_rate = rates.get("QA (авто)", rates.get("QA", 15))
        qa_client_rate = qa_rate * (1 + margin / 100)
        specialists.append({
            "name": "QA",
            "weeks": round(qa_weeks, 2),
            "days": round(qa_weeks * 5, 1),
            "hours": round(qa_hours, 1),
            "rate": qa_rate,
            "cost": round(qa_hours * qa_rate, 2),
            "client_rate": round(qa_client_rate, 2),
            "client_cost": round(qa_hours * qa_client_rate, 2),
        })

    session["specialists_final"] = specialists

    # Extract project name from spec
    spec_lines = session.get("spec", "").split("\n")
    project_name = "Проект"
    for line in spec_lines:
        clean = line.strip().strip("#").strip("*").strip()
        if clean:
            project_name = clean[:80]
            break

    total_cost = sum(s["cost"] for s in specialists)
    total_client = sum(s["client_cost"] for s in specialists)

    # Show summary in chat
    lines = ["📊 ИТОГ ОЦЕНКИ", "=" * 50, ""]
    lines.append(f"Коэффициент K: {K:.2f}")
    lines.append(f"Маржа: {margin:.0f}%")
    lines.append("")
    for s in specialists:
        lines.append(
            f"• {s['name']}: {s['hours']}ч | ${s['rate']}/ч → ${s['client_rate']:.0f}/ч | "
            f"Себестоимость ${s['cost']:,.0f} → Клиент ${s['client_cost']:,.0f}"
        )
    lines.append("")
    lines.append(f"Себестоимость: ${total_cost:,.0f}")
    lines.append(f"Стоимость для клиента: ${total_client:,.0f}")
    lines.append(f"Маржинальность: {(1 - total_cost/total_client)*100:.1f}%")

    for chunk in _split_message("\n".join(lines)):
        await update.message.reply_text(chunk)

    # Generate xlsx
    xlsx_bytes = _build_report_xlsx(
        project_name, specialists, session["modules"], margin, coeffs, K
    )

    await update.message.reply_document(
        document=xlsx_bytes,
        filename="Оценка_проекта.xlsx",
        caption="📊 Оценка проекта",
    )

    await update.message.reply_text(
        "✅ Оценка проекта завершена!\n\nДля нового запроса: /new_request"
    )
    return ConversationHandler.END


def _build_report_xlsx(
    project_name: str,
    specialists: list[dict],
    modules: list[dict],
    margin_pct: float,
    coeffs: dict,
    K: float,
) -> io.BytesIO:
    """Build client-facing xlsx report matching the exact template design.

    specialists: list of dicts with keys name, weeks, days, hours, rate, cost, client_rate, client_cost
    modules: list of dicts with keys name, tasks (each task has task, specialist, comment, min_days, max_days)
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "Для клиента"

    # ── Fonts (Montserrat as in template) ─────────────────────────────────
    font_title = Font(name="Montserrat", size=24)
    font_subtitle = Font(name="Montserrat", size=10)
    font_date = Font(name="Montserrat", size=10, color="777777")
    font_header = Font(name="Montserrat", size=12)
    font_normal = Font(name="Montserrat", size=10)
    font_section = Font(name="Montserrat", size=11, bold=True)

    # ── Fills ─────────────────────────────────────────────────────────────
    fill_orange = PatternFill(start_color="FFA301", end_color="FFA301", fill_type="solid")
    fill_gray = PatternFill(start_color="F4F4F4", end_color="F4F4F4", fill_type="solid")
    fill_white = PatternFill(start_color="FFFFFF", end_color="FFFFFF", fill_type="solid")

    def _apply_outer_border(ws, start_row, end_row, start_col, end_col):
        """Apply thin outer border to a rectangular range."""
        for r in range(start_row, end_row + 1):
            for c in range(start_col, end_col + 1):
                cell = ws.cell(row=r, column=c)
                top = Side("thin") if r == start_row else None
                bottom = Side("thin") if r == end_row else None
                left = Side("thin") if c == start_col else None
                right = Side("thin") if c == end_col else None
                cell.border = Border(top=top, bottom=bottom, left=left, right=right)

    align_center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    align_left = Alignment(horizontal="left", vertical="center", wrap_text=True)

    # ── Column widths (match template) ────────────────────────────────────
    ws.column_dimensions["A"].width = 7.8
    ws.column_dimensions["B"].width = 21
    ws.column_dimensions["C"].width = 39.5
    ws.column_dimensions["D"].width = 35
    ws.column_dimensions["E"].width = 27

    # ── Logo (top-left) ─────────────────────────────────────────────────
    logo_path = os.path.join(os.path.dirname(__file__), "metalamp-logo.png")
    if os.path.exists(logo_path):
        img = XlImage(logo_path)
        orig_w, orig_h = img.width, img.height
        target_w = 300
        img.width = target_w
        img.height = int(orig_h * (target_w / max(orig_w, 1)))
        ws.add_image(img, "B2")

    # ── Row 5: Title ──────────────────────────────────────────────────────
    row = 5
    ws.row_dimensions[row].height = 38.25
    ws.merge_cells(f"B{row}:C{row}")
    ws.cell(row=row, column=2, value="Оценка проекта ").font = font_title
    ws.merge_cells(f"D{row}:E{row}")
    ws.cell(row=row, column=4, value=project_name).font = font_title

    # ── Row 6: Subtitle 1 ────────────────────────────────────────────────
    row = 6
    ws.row_dimensions[row].height = 38.25
    ws.merge_cells(f"B{row}:E{row}")
    ws.cell(row=row, column=2, value="В стоимость входит тестирование, работа менеджера").font = font_subtitle

    # ── Row 7: Subtitle 2 ────────────────────────────────────────────────
    row = 7
    ws.row_dimensions[row].height = 31.5
    ws.merge_cells(f"B{row}:E{row}")
    ws.cell(row=row, column=2, value="В течение 3 месяцев мы бесплатно устраняем технические ошибки (техподдержка)").font = font_subtitle

    # ── Row 8: Date ───────────────────────────────────────────────────────
    row = 8
    ws.cell(row=row, column=2, value=f"Актуально на: {date.today().strftime('%d.%m.%Y')}").font = font_date

    # ── Row 9-10: Summary block ───────────────────────────────────────────
    total_hours = sum(s["hours"] for s in specialists)
    total_client_cost = sum(s["client_cost"] for s in specialists)
    team_size = len(specialists)

    row = 9
    ws.row_dimensions[row].height = 30
    for col, val in [(2, "Команда проекта,\nчеловек"), (3, "Длительность проекта,\nчасы"), (4, "Стоимость, $")]:
        c = ws.cell(row=row, column=col, value=val)
        c.font = font_header
        c.fill = fill_orange
        c.alignment = align_center
    _apply_outer_border(ws, 9, 9, 2, 4)

    row = 10
    ws.row_dimensions[row].height = 15
    for col, val in [(2, team_size), (3, round(total_hours)), (4, round(total_client_cost))]:
        c = ws.cell(row=row, column=col, value=val)
        c.font = font_normal
        c.alignment = align_center
    _apply_outer_border(ws, 10, 10, 2, 4)

    # ── Row 12+: Specialists table ────────────────────────────────────────
    row = 12
    ws.row_dimensions[row].height = 15
    for col, val in [(2, "Специалисты"), (3, "Занятость на проекте, недели"), (4, "Занятость на проекте, часы")]:
        c = ws.cell(row=row, column=col, value=val)
        c.font = font_header
        c.fill = fill_orange
        c.alignment = align_center
    _apply_outer_border(ws, 12, 12, 2, 4)

    spec_start_row = row + 1
    for idx, s in enumerate(specialists):
        row = spec_start_row + idx
        is_odd = idx % 2 == 0
        fill = fill_gray if is_odd else fill_white
        for col, val in [(2, s["name"]), (3, s["weeks"]), (4, round(s["hours"]))]:
            c = ws.cell(row=row, column=col, value=val)
            c.font = font_normal
            c.fill = fill
            c.alignment = align_center if col >= 3 else align_left
        ws.cell(row=row, column=2).border = Border(left=Side("thin"))
        ws.cell(row=row, column=4).border = Border(right=Side("thin"))

    last_spec_row = spec_start_row + len(specialists) - 1
    for c in range(2, 5):
        cell = ws.cell(row=last_spec_row, column=c)
        existing = cell.border
        cell.border = Border(
            left=existing.left, right=existing.right,
            top=existing.top, bottom=Side("thin"),
        )

    row = last_spec_row + 3

    # ── Task decomposition by modules ─────────────────────────────────────
    for module in modules:
        # Module title (bold, no fill)
        ws.row_dimensions[row].height = 13.8
        ws.cell(row=row, column=2, value=module["name"]).font = font_section
        row += 1

        # Orange header row
        ws.row_dimensions[row].height = 15
        for col, val in [(2, "Специалист"), (3, "Задача"), (4, "Комментарий"), (5, "Оценка, дни")]:
            c = ws.cell(row=row, column=col, value=val)
            c.font = font_header
            c.fill = fill_orange
            c.alignment = align_center
        _apply_outer_border(ws, row, row, 2, 5)
        row += 1

        # Task rows
        for idx, t in enumerate(module.get("tasks", [])):
            is_odd = idx % 2 == 0
            fill = fill_gray if is_odd else fill_white
            days_str = f"{t['min_days']}-{t['max_days']}"
            comment = t.get("comment", "")
            for col, val in [(2, t["specialist"]), (3, t["task"]), (4, comment), (5, days_str)]:
                c = ws.cell(row=row, column=col, value=val)
                c.font = font_normal
                c.fill = fill
                c.alignment = align_center if col == 5 else align_left
            ws.cell(row=row, column=2).border = Border(left=Side("thin"))
            ws.cell(row=row, column=5).border = Border(right=Side("thin"))
            row += 1

        # Bottom border on last task row
        if module.get("tasks"):
            last_task_row = row - 1
            for c in range(2, 6):
                cell = ws.cell(row=last_task_row, column=c)
                existing = cell.border
                cell.border = Border(
                    left=existing.left, right=existing.right,
                    top=existing.top, bottom=Side("thin"),
                )

        row += 1  # gap between module sections

    # ── Sheet 2: Оценка (decomposition matching P2P reference) ──────────
    _build_estimation_sheet(wb, modules, specialists, coeffs, K)

    # ── Sheet 3: GANTT Chart ──────────────────────────────────────────────
    _build_gantt_sheet(wb, modules, K)

    # Save
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    buf.name = "report.xlsx"
    return buf


def _build_estimation_sheet(
    wb: Workbook,
    modules: list[dict],
    specialists: list[dict],
    coeffs: dict,
    K: float,
):
    """Add 'Оценка' sheet matching the P2P reference format."""
    ws = wb.create_sheet("Оценка")

    font_bold = Font(name="Arial", bold=True)
    font_normal = Font(name="Arial")
    font_header = Font(name="Arial", bold=True)
    font_module = Font(name="Arial", size=11, bold=True)
    font_task = Font(name="Arial", size=11)

    align_center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    align_left = Alignment(horizontal="left", vertical="center", wrap_text=True)
    align_right = Alignment(horizontal="right", vertical="center")

    # Column widths (match reference)
    ws.column_dimensions["A"].width = 10
    ws.column_dimensions["B"].width = 19
    ws.column_dimensions["C"].width = 25
    ws.column_dimensions["D"].width = 29
    ws.column_dimensions["E"].width = 10
    ws.column_dimensions["F"].width = 13
    ws.column_dimensions["G"].width = 13
    ws.column_dimensions["H"].width = 15

    # ── Rows 1-3: Instruction ─────────────────────────────────────────
    ws.cell(row=2, column=2, value=(
        "Инструкция:\n"
        "1. Задачи оцениваются в днях/полу днях.\n"
        "2. Столбец B — вид работ (специалист).\n"
        "3. Если задача > 5 дней — разбить на подзадачи.\n"
        "4. Задачи по фронтенду включают интеграцию с бэкендом.\n"
        "5. Столбцы F/G — первичная оценка, H — с учётом коэффициента K."
    )).font = font_normal
    ws.cell(row=2, column=2).alignment = Alignment(wrap_text=True, vertical="top")
    ws.merge_cells("B2:D3")

    # ── Row 4: Summary headers ────────────────────────────────────────
    row = 4
    for col, val in [(2, "Дни минимум"), (3, "Дни максимум"), (4, f"Недель с коф. (K={K:.2f})")]:
        c = ws.cell(row=row, column=col, value=val)
        c.font = font_bold
        c.alignment = align_center if col >= 3 else align_left

    # ── Rows 5+: Summary per specialist ───────────────────────────────
    # Collect min/max days per specialist from modules
    spec_summary: dict[str, dict] = {}
    for m in modules:
        for t in m.get("tasks", []):
            name = t["specialist"]
            if name not in spec_summary:
                spec_summary[name] = {"min": 0, "max": 0}
            spec_summary[name]["min"] += t.get("min_days", 0)
            spec_summary[name]["max"] += t.get("max_days", 0)

    # Add PM/QA from specialists list if they're auto-calculated
    for s in specialists:
        if s["name"] not in spec_summary:
            # Auto-calculated specialist (PM, QA auto)
            spec_summary[s["name"]] = {
                "min": s["days"],
                "max": s["days"],
            }

    summary_start = 5
    spec_num = 0
    spec_number_map = {}  # specialist name -> number
    for idx, (name, d) in enumerate(spec_summary.items()):
        r = summary_start + idx
        spec_num = idx + 1
        spec_number_map[name] = spec_num
        avg = (d["min"] + d["max"]) / 2
        weeks_k = round(avg * K / 5, 2)
        ws.cell(row=r, column=1, value=spec_num).font = font_normal
        ws.cell(row=r, column=1).alignment = align_center
        ws.cell(row=r, column=2, value=d["min"]).font = font_normal
        ws.cell(row=r, column=2).alignment = align_right
        ws.cell(row=r, column=3, value=d["max"]).font = font_normal
        ws.cell(row=r, column=3).alignment = align_right
        ws.cell(row=r, column=4, value=weeks_k).font = font_bold
        ws.cell(row=r, column=4).alignment = align_right
        # Specialist name label in col E
        ws.cell(row=r, column=5, value=name).font = font_normal

    # ── Header row for task table ─────────────────────────────────────
    row = summary_start + len(spec_summary) + 1
    headers = {
        1: "Распределение работ",
        2: "Вид работ",
        3: "Задача",
        4: "Комментарий",
        6: "Минимальная оценка дни",
        7: "Максимальная оценка дни",
        8: "Итого с коэф.",
    }
    for col, val in headers.items():
        c = ws.cell(row=row, column=col, value=val)
        c.font = font_header
        c.alignment = align_center
    row += 1

    # ── Module headers and task rows ──────────────────────────────────
    for module in modules:
        # Module header row — merged C:D, bold
        ws.merge_cells(f"C{row}:D{row}")
        c = ws.cell(row=row, column=3, value=module["name"])
        c.font = font_module
        c.alignment = align_center
        row += 1

        for t in module.get("tasks", []):
            spec_name = t["specialist"]
            spec_n = spec_number_map.get(spec_name, "")
            min_d = t.get("min_days", 0)
            max_d = t.get("max_days", 0)
            avg_d = (min_d + max_d) / 2
            final_d = round(avg_d * K, 1)

            ws.cell(row=row, column=1, value=spec_n).font = font_normal
            ws.cell(row=row, column=1).alignment = align_right

            c = ws.cell(row=row, column=2, value=spec_name)
            c.font = font_bold
            c.alignment = align_left

            ws.cell(row=row, column=3, value=t["task"]).font = font_task
            ws.cell(row=row, column=3).alignment = align_left

            ws.cell(row=row, column=4, value=t.get("comment", "")).font = font_task
            ws.cell(row=row, column=4).alignment = align_left

            ws.cell(row=row, column=6, value=min_d).font = font_normal
            ws.cell(row=row, column=6).alignment = align_center

            ws.cell(row=row, column=7, value=max_d).font = font_normal
            ws.cell(row=row, column=7).alignment = align_center

            ws.cell(row=row, column=8, value=final_d).font = font_normal
            ws.cell(row=row, column=8).alignment = align_center

            row += 1


def _build_gantt_sheet(wb: Workbook, modules: list[dict], K: float):
    """Add 'GANTT Chart' sheet with colored bars per specialist/phase."""

    # ── Workday helpers ───────────────────────────────────────────────
    def next_workday(dt):
        while dt.weekday() >= 5:
            dt += timedelta(days=1)
        return dt

    def add_workdays(start_dt, num_workdays):
        cur = next_workday(start_dt)
        counted = 1
        while counted < num_workdays:
            cur += timedelta(days=1)
            if cur.weekday() < 5:
                counted += 1
        return cur

    # ── Phase colors ──────────────────────────────────────────────────
    PHASE_COLORS = [
        {"header": "1F4E79", "fill": "D6E4F0", "bar": "5B9BD5"},
        {"header": "7B2D26", "fill": "F2DCDB", "bar": "C0504D"},
        {"header": "4F6228", "fill": "EBF1DE", "bar": "9BBB59"},
        {"header": "31859C", "fill": "DAEEF3", "bar": "4BACC6"},
        {"header": "E36C09", "fill": "FDE9D9", "bar": "F79646"},
        {"header": "60497A", "fill": "E4DFEC", "bar": "8064A2"},
        {"header": "4A452A", "fill": "F2F2E6", "bar": "948A54"},
    ]

    month_names_ru = {
        1: "Январь", 2: "Февраль", 3: "Март", 4: "Апрель",
        5: "Май", 6: "Июнь", 7: "Июль", 8: "Август",
        9: "Сентябрь", 10: "Октябрь", 11: "Ноябрь", 12: "Декабрь",
    }

    DATA_COL_START = 7  # column G

    # ── Build flat task list grouped by specialist (phase) ────────────
    # Each specialist becomes a "phase" with their own color
    specialist_tasks: dict[str, list[dict]] = {}
    for module in modules:
        for t in module.get("tasks", []):
            spec = t["specialist"]
            if spec not in specialist_tasks:
                specialist_tasks[spec] = []
            min_d = t.get("min_days", 0)
            max_d = t.get("max_days", 0)
            avg_d = (min_d + max_d) / 2
            duration_days = max(1, round(avg_d * K))
            specialist_tasks[spec].append({
                "task": t["task"],
                "duration": duration_days,
            })

    if not specialist_tasks:
        return

    # Assign colors to specialists
    spec_list = list(specialist_tasks.keys())
    spec_colors = {}
    for i, spec in enumerate(spec_list):
        spec_colors[spec] = PHASE_COLORS[i % len(PHASE_COLORS)]

    # ── Schedule tasks with cascade logic ─────────────────────────────
    # All phases start from today (next workday)
    project_start_raw = next_workday(date.today())
    # Convert to datetime for consistency
    project_start_dt = datetime(project_start_raw.year, project_start_raw.month, project_start_raw.day)

    scheduled_tasks = []  # (specialist, task_name, start_dt, end_dt, duration)
    for spec in spec_list:
        current_start = project_start_dt
        for t in specialist_tasks[spec]:
            start = next_workday(current_start)
            end = add_workdays(start, t["duration"])
            scheduled_tasks.append((spec, t["task"], start, end, t["duration"]))
            # Next task starts after this one ends
            current_start = end + timedelta(days=1)

    if not scheduled_tasks:
        return

    # ── Calendar range ────────────────────────────────────────────────
    project_start = min(t[2] for t in scheduled_tasks)
    while project_start.weekday() != 0:  # Monday
        project_start -= timedelta(days=1)
    project_end = max(t[3] for t in scheduled_tasks)
    while project_end.weekday() != 4:  # Friday
        project_end += timedelta(days=1)

    all_days = []
    d = project_start
    while d <= project_end:
        all_days.append(d)
        d += timedelta(days=1)
    num_days = len(all_days)

    # ── Create sheet ──────────────────────────────────────────────────
    ws = wb.create_sheet("GANTT Chart")

    # ── Styles ────────────────────────────────────────────────────────
    header_fill = PatternFill(start_color="2F5496", end_color="2F5496", fill_type="solid")
    weekend_header_fill = PatternFill(start_color="C0C0C0", end_color="C0C0C0", fill_type="solid")
    date_label_fill = PatternFill(start_color="D6DCE4", end_color="D6DCE4", fill_type="solid")

    thin_border = Border(
        left=Side(style="thin", color="BFBFBF"),
        right=Side(style="thin", color="BFBFBF"),
        top=Side(style="thin", color="BFBFBF"),
        bottom=Side(style="thin", color="BFBFBF"),
    )
    week_sep_border = Border(
        left=Side(style="thin", color="BFBFBF"),
        right=Side(style="medium", color="808080"),
        top=Side(style="thin", color="BFBFBF"),
        bottom=Side(style="thin", color="BFBFBF"),
    )

    header_font = Font(name="Calibri", size=10, bold=True, color="FFFFFF")
    month_font = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
    day_num_font = Font(name="Calibri", size=7, color="44546A")
    weekday_font = Font(name="Calibri", size=6, color="808080")
    weekday_bold_font = Font(name="Calibri", size=6, bold=True, color="999999")
    task_font = Font(name="Calibri", size=9)
    phase_header_font = Font(name="Calibri", size=10, bold=True, color="FFFFFF")

    center_align = Alignment(horizontal="center", vertical="center")
    left_wrap = Alignment(vertical="center", wrap_text=True)
    phase_align = Alignment(horizontal="left", vertical="center")

    weekday_names = ["Пн", "Вт", "Ср", "Чт", "Пт", "Сб", "Вс"]

    # ── Column widths ─────────────────────────────────────────────────
    ws.column_dimensions["A"].width = 22
    ws.column_dimensions["B"].width = 14
    ws.column_dimensions["C"].width = 42
    ws.column_dimensions["D"].width = 11
    ws.column_dimensions["E"].width = 6
    ws.column_dimensions["F"].width = 11
    for i in range(num_days):
        col_letter = get_column_letter(DATA_COL_START + i)
        ws.column_dimensions[col_letter].width = 2.5 if all_days[i].weekday() >= 5 else 3.8

    def _border_for_day(day):
        return week_sep_border if day.weekday() == 6 else thin_border

    # ── Row 1: Months ─────────────────────────────────────────────────
    ws.row_dimensions[1].height = 20
    for c in range(1, 7):
        cell = ws.cell(row=1, column=c)
        cell.fill = header_fill
        cell.border = thin_border

    # Find month spans
    month_spans = []
    if all_days:
        cur_month = (all_days[0].year, all_days[0].month)
        span_start = 0
        for i, day in enumerate(all_days):
            m = (day.year, day.month)
            if m != cur_month:
                month_spans.append((cur_month, span_start, i - 1))
                cur_month = m
                span_start = i
        month_spans.append((cur_month, span_start, len(all_days) - 1))

    for (year, month), start_idx, end_idx in month_spans:
        start_col = DATA_COL_START + start_idx
        end_col = DATA_COL_START + end_idx
        if end_col > start_col:
            ws.merge_cells(
                start_row=1, start_column=start_col,
                end_row=1, end_column=end_col,
            )
        cell = ws.cell(row=1, column=start_col)
        cell.value = f"{month_names_ru[month]} {year}"
        cell.font = month_font
        cell.fill = header_fill
        cell.alignment = center_align
        for ci in range(start_col, end_col + 1):
            c = ws.cell(row=1, column=ci)
            c.fill = header_fill
            c.border = _border_for_day(all_days[ci - DATA_COL_START])

    # ── Row 2: Column headers + day numbers ───────────────────────────
    ws.row_dimensions[2].height = 18
    col_headers = ["Фаза", "Роль", "Задача", "Старт", "Дней", "Конец"]
    for i, val in enumerate(col_headers):
        cell = ws.cell(row=2, column=i + 1, value=val)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = center_align
        cell.border = thin_border

    for i, day in enumerate(all_days):
        col = DATA_COL_START + i
        cell = ws.cell(row=2, column=col, value=day.day)
        cell.font = day_num_font
        cell.fill = weekend_header_fill if day.weekday() >= 5 else date_label_fill
        cell.alignment = center_align
        cell.border = _border_for_day(day)

    # ── Row 3: Weekday names ──────────────────────────────────────────
    ws.row_dimensions[3].height = 14
    for c in range(1, 7):
        cell = ws.cell(row=3, column=c)
        cell.fill = date_label_fill
        cell.border = thin_border

    for i, day in enumerate(all_days):
        col = DATA_COL_START + i
        cell = ws.cell(row=3, column=col, value=weekday_names[day.weekday()])
        cell.font = weekday_bold_font if day.weekday() >= 5 else weekday_font
        cell.fill = weekend_header_fill if day.weekday() >= 5 else date_label_fill
        cell.alignment = center_align
        cell.border = _border_for_day(day)

    # ── Data rows ─────────────────────────────────────────────────────
    row = 4
    current_spec = None

    for spec, task_name, start_dt, end_dt, duration in scheduled_tasks:
        colors = spec_colors[spec]

        # Weekend row color (fill darkened by 25)
        rb = int(colors["fill"][:2], 16)
        gb = int(colors["fill"][2:4], 16)
        bb = int(colors["fill"][4:6], 16)
        wknd_hex = f"{max(0, rb - 25):02X}{max(0, gb - 25):02X}{max(0, bb - 25):02X}"

        fill_phase = PatternFill(start_color=colors["fill"], end_color=colors["fill"], fill_type="solid")
        fill_bar = PatternFill(start_color=colors["bar"], end_color=colors["bar"], fill_type="solid")
        fill_wknd = PatternFill(start_color=wknd_hex, end_color=wknd_hex, fill_type="solid")
        fill_header = PatternFill(start_color=colors["header"], end_color=colors["header"], fill_type="solid")

        # Phase header row (when specialist changes)
        if spec != current_spec:
            current_spec = spec
            ws.row_dimensions[row].height = 22
            ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=6)
            cell = ws.cell(row=row, column=1, value=spec)
            cell.font = phase_header_font
            cell.fill = fill_header
            cell.alignment = phase_align
            # Fill all columns including data area
            for c in range(1, 7):
                cl = ws.cell(row=row, column=c)
                cl.fill = fill_header
                cl.border = thin_border
            for i, day in enumerate(all_days):
                col = DATA_COL_START + i
                cl = ws.cell(row=row, column=col)
                cl.fill = fill_header
                cl.border = _border_for_day(day)
            row += 1

        # Task row
        ws.row_dimensions[row].height = 20
        task_data = [
            (1, spec, left_wrap),
            (2, spec, left_wrap),
            (3, task_name, left_wrap),
            (4, start_dt.strftime("%d.%m.%y"), center_align),
            (5, duration, center_align),
            (6, end_dt.strftime("%d.%m.%y"), center_align),
        ]
        for col, val, align in task_data:
            cell = ws.cell(row=row, column=col, value=val)
            cell.font = task_font
            cell.fill = fill_phase
            cell.alignment = align
            cell.border = thin_border

        # Data area — bars
        for i, day in enumerate(all_days):
            col = DATA_COL_START + i
            cell = ws.cell(row=row, column=col)
            cell.border = _border_for_day(day)

            if start_dt <= day <= end_dt and day.weekday() < 5:
                cell.fill = fill_bar
            elif day.weekday() >= 5:
                cell.fill = fill_wknd
            else:
                cell.fill = fill_phase

        row += 1

    # ── Freeze panes ──────────────────────────────────────────────────
    ws.freeze_panes = "G4"


async def cancel(update: Update, context) -> int:
    await update.message.reply_text("Оценка отменена. Для нового запроса: /new_request")
    return ConversationHandler.END


async def skip(update: Update, context) -> int:
    user_id = update.effective_user.id
    session = get_session(user_id)
    # Cancel background task if running
    session["_cancelled"] = True
    task = session.pop("_task", None)
    if task and not task.done():
        task.cancel()
    await update.message.reply_text("⏹ Анализ прерван. Для нового запроса: /new_request")
    return ConversationHandler.END


async def handle_processing_callback(update: Update, context) -> int:
    """Handle callbacks that arrive while in PROCESSING state (after background task finishes)."""
    query = update.callback_query
    await query.answer()
    user_id = query.from_user.id
    session = get_session(user_id)

    next_state = session.pop("_proc_next_state", None)

    # Route to the right handler based on what the background task prepared
    if query.data == "decompose":
        return await _decompose(query, session)
    elif query.data == "redo_spec":
        await query.message.reply_text(
            "Опиши, что не так в текущем ТЗ. Какие замечания нужно учесть?"
        )
        return WAITING_SPEC_FEEDBACK
    elif query.data == "skip_answers":
        return await _generate_spec_from_msg(query.message, session)

    return next_state if next_state is not None else PROCESSING


async def handle_processing_text(update: Update, context) -> int:
    """Handle text messages while in PROCESSING state."""
    user_id = update.effective_user.id
    session = get_session(user_id)
    next_state = session.get("_proc_next_state")

    if next_state == WAITING_ANSWERS:
        # Background task finished and asked questions — this is the answer
        session["answers"] = update.message.text
        return await _generate_spec(update, session)

    await update.message.reply_text("⏳ Обработка ещё идёт... Для отмены: /skip")
    return PROCESSING


def _split_message(text: str, limit: int = 4000) -> list[str]:
    """Split long text into chunks for Telegram."""
    if len(text) <= limit:
        return [text]
    chunks = []
    while text:
        if len(text) <= limit:
            chunks.append(text)
            break
        split_at = text.rfind("\n", 0, limit)
        if split_at == -1:
            split_at = limit
        chunks.append(text[:split_at])
        text = text[split_at:].lstrip("\n")
    return chunks


def main():
    global claude_client

    bot_token = os.getenv("TELEGRAM_BOT_TOKEN")
    anthropic_key = os.getenv("ANTHROPIC_API_KEY")

    if not bot_token:
        raise RuntimeError("TELEGRAM_BOT_TOKEN not set in .env")
    if not anthropic_key:
        raise RuntimeError("ANTHROPIC_API_KEY not set in .env")

    claude_client = claude_service.create_client(anthropic_key)

    app = Application.builder().token(bot_token).build()

    conv_handler = ConversationHandler(
        entry_points=[CommandHandler("new_request", cmd_new_request)],
        states={
            WAITING_REQUEST: [
                MessageHandler(
                    (filters.TEXT | filters.Document.ALL) & ~filters.COMMAND,
                    handle_request,
                )
            ],
            WAITING_ANSWERS: [
                CallbackQueryHandler(handle_skip_answers, pattern="^skip_answers$"),
                MessageHandler(filters.TEXT & ~filters.COMMAND, handle_answers),
            ],
            REVIEWING_SPEC: [
                CallbackQueryHandler(handle_spec_review, pattern="^(decompose|redo_spec)$")
            ],
            WAITING_SPEC_FEEDBACK: [
                MessageHandler(filters.TEXT & ~filters.COMMAND, handle_spec_feedback)
            ],
            WAITING_COEFFICIENTS: [
                CallbackQueryHandler(handle_coefficients_done, pattern="^coeffs_done$"),
                MessageHandler(filters.TEXT & ~filters.COMMAND, handle_coefficients),
            ],
            WAITING_RATES: [
                CallbackQueryHandler(handle_rates_done, pattern="^rates_done$"),
                MessageHandler(filters.TEXT & ~filters.COMMAND, handle_rates),
            ],
            WAITING_MARGIN: [
                MessageHandler(filters.TEXT & ~filters.COMMAND, handle_margin)
            ],
            PROCESSING: [
                CallbackQueryHandler(
                    handle_processing_callback,
                    pattern="^(decompose|redo_spec|skip_answers)$",
                ),
                MessageHandler(filters.TEXT & ~filters.COMMAND, handle_processing_text),
            ],
        },
        fallbacks=[
            CommandHandler("skip", skip),
            CommandHandler("cancel", cancel),
        ],
    )

    app.add_handler(CommandHandler("start", cmd_start))
    app.add_handler(conv_handler)

    logger.info("Bot started")
    app.run_polling(allowed_updates=Update.ALL_TYPES)


if __name__ == "__main__":
    main()
