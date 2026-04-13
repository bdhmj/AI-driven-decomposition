"""Build a .docx specification from markdown text.

Usage:
    python scripts/build_docx.py input/spec.md output/Техническое_задание.docx
"""

import io
import re
import sys
from docx import Document
from docx.shared import Pt


def build_spec_docx(spec_text: str) -> io.BytesIO:
    """Convert markdown spec text (with headings and tables) into a formatted .docx."""
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)

    lines = spec_text.split("\n")
    i = 0
    while i < len(lines):
        stripped = lines[i].strip()

        if not stripped:
            doc.add_paragraph("")
            i += 1
            continue

        # Detect markdown table (starts with |)
        if stripped.startswith("|") and "|" in stripped[1:]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            rows_data = []
            for tl in table_lines:
                if re.match(r"^\|[\s\-:|]+\|$", tl):
                    continue
                cells = [c.strip() for c in tl.split("|")[1:-1]]
                rows_data.append(cells)
            if rows_data:
                num_cols = max(len(r) for r in rows_data)
                table = doc.add_table(rows=len(rows_data), cols=num_cols)
                table.style = "Light Grid Accent 1"
                for ri, row_cells in enumerate(rows_data):
                    for ci, cell_val in enumerate(row_cells):
                        if ci < num_cols:
                            cell = table.rows[ri].cells[ci]
                            cell.text = re.sub(r"\*\*(.*?)\*\*", r"\1", cell_val)
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.name = "Arial"
                                    run.font.size = Pt(9)
                    if ri == 0:
                        for ci in range(num_cols):
                            if ci < len(table.rows[0].cells):
                                for run in table.rows[0].cells[ci].paragraphs[0].runs:
                                    run.bold = True
                doc.add_paragraph("")
            continue

        # Markdown headings
        heading_match = re.match(r"^(#{1,4})\s+(.*)", stripped)
        if heading_match:
            level = min(len(heading_match.group(1)), 4)
            text = heading_match.group(2).strip("*").strip()
            doc.add_heading(text, level=level)
            i += 1
            continue

        # Numbered section titles (e.g. "1. Title" or "1.2. Title")
        section_match = re.match(r"^(\d+\.[\d.]*)\s+\*\*(.*?)\*\*", stripped)
        if section_match:
            doc.add_heading(f"{section_match.group(1)} {section_match.group(2)}", level=2)
            i += 1
            continue

        # Bold lines (full line wrapped in **)
        if stripped.startswith("**") and stripped.endswith("**") and len(stripped) > 4:
            p = doc.add_paragraph()
            run = p.add_run(re.sub(r"^\*\*(.*)\*\*$", r"\1", stripped))
            run.bold = True
            i += 1
            continue

        # Bullet points (- or •, with or without space)
        bullet_match = re.match(r"^[-•]\s?(.*)", stripped)
        if bullet_match:
            doc.add_paragraph(bullet_match.group(1), style="List Bullet")
            i += 1
            continue

        # Regular paragraph
        clean = re.sub(r"\*\*(.*?)\*\*", r"\1", stripped)
        doc.add_paragraph(clean)
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    buf.name = "spec.docx"
    return buf


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python scripts/build_docx.py <input.md> <output.docx>")
        sys.exit(1)

    with open(sys.argv[1], "r", encoding="utf-8") as f:
        spec_text = f.read()

    result = build_spec_docx(spec_text)
    with open(sys.argv[2], "wb") as f:
        f.write(result.read())

    print(f"Saved: {sys.argv[2]}")
